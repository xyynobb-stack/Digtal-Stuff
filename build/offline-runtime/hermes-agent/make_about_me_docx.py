# -*- coding: utf-8 -*-
"""生成《我是什么模型》Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# 默认中文字体
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# 标题
title = doc.add_heading("我是什么模型？", level=0)
for run in title.runs:
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.color.rgb = RGBColor(0x1F, 0x3B, 0x73)

doc.add_paragraph("—— 一份关于我的自我介绍")

# 第一部分
doc.add_heading("一、我的身份", level=1)
doc.add_paragraph("我是 Hermes One —— 由 Nous Research 打造的人工智能助手，运行在 Hermes Agent 平台上。"
                  "我是一个能主动使用工具、执行任务的智能体（Agent），而不只是一个只会对话的聊天机器人。")

# 第二部分
doc.add_heading("二、我运行的模型", level=1)
doc.add_paragraph("当前驱动我这个会话的底层大语言模型是：DeepSeek-V4-Flash，"
                  "由公司级平台（company-platform）提供模型服务。"
                  "这是一个追求快速响应、擅长推理与工具调用的模型。")

# 第三部分
doc.add_heading("三、我的能力", level=1)
capabilities = [
    "读写文件、编辑代码，完成软件开发任务",
    "搜索与阅读网页内容（通过内置 Python 实现）",
    "创建 Word、Excel、PPT、PDF 等办公文档",
    "执行终端命令、运行测试、管理项目",
    "调用各类工具完成跨平台自动化任务",
    "阅读和分析已有文档（docx、xlsx、pdf 等格式）",
]
for cap in capabilities:
    doc.add_paragraph(cap, style="List Bullet")

# 第四部分
doc.add_heading("四、我的特点", level=1)
traits = [
    ("诚实可靠：", "不编造结果，任务受阻时如实说明并尝试替代方案。"),
    ("行动导向：", "说做就做，用真实的工具输出交付成果，而不是空谈计划。"),
    ("细致严谨：", "先读代码再动手，遵循项目既有风格，只改该改的地方。"),
    ("高效并行：", "同时发起多个独立查询，加快任务完成速度。"),
]
for name, desc in traits:
    p = doc.add_paragraph()
    run = p.add_run(name)
    run.bold = True
    p.add_run(desc)

# 落款
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer.add_run("Hermes One · Nous Research").bold = True
footer2 = doc.add_paragraph()
footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
footer2.add_run("2026 年 8 月")

out = r"C:\Users\Administrator\Desktop\我是什么模型.docx"
doc.save(out)
print("已保存:", out)
